from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt, RGBColor, Inches
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml import OxmlElement as OE
import win32com.client
from Contenido import *
import pandas as pd
import numpy as np

# Create document
document = Document("base.docx")


def crear_tabla(df):
    table = document.add_table(rows=1, cols=len(df.columns))
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    for c, value in enumerate(df.columns):
        hdr_cells[c].text = value
        hdr_cells[c].paragraphs[0].runs[0].font.size = Pt(14)
        hdr_cells[c].paragraphs[0].runs[0].font.bold = True

    for n in range(len(df.columns)):
        # GET CELLS XML ELEMENT
        cell_xml_element = table.rows[0].cells[n]._tc
        # RETRIEVE THE TABLE CELL PROPERTIES
        table_cell_properties = cell_xml_element.get_or_add_tcPr()
        # CREATE SHADING OBJECT
        shade_obj = OxmlElement('w:shd')
        # SET THE SHADING OBJECT
        shade_obj.set(qn('w:fill'), "35D59C")
        # APPEND THE PROPERTIES TO THE TABLE CELL PROPERTIES
        table_cell_properties.append(shade_obj)

    for index, row in df.iterrows():
        row_cells = table.add_row().cells
        for n in range(len(df.columns)):
            row_cells[n].text = str(row[df.columns[n]])
def add_list_of_table():
    paragraph = document.add_paragraph()
    run = paragraph.add_run()
    fldChar = OE('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'begin')
    fldChar.set(qn('w:dirty'), 'true')
    instrText = OE('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'TOC \\h \\z \\c "Table"'  # "Table" of list of table and "Figure" for list of figure
    fldChar2 = OE('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    fldChar3 = OE('w:t')
    fldChar3.text = "Right-click to update field."
    fldChar2.append(fldChar3)

    fldChar4 = OE('w:fldChar')
    fldChar4.set(qn('w:fldCharType'), 'end')

    run._r.append(fldChar)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run._r.append(fldChar4)
def add_list_of_figure():
    paragraph = document.add_paragraph()
    run = paragraph.add_run()
    fldChar = OE('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'begin')
    fldChar.set(qn('w:dirty'), 'true')
    instrText = OE('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'TOC \\h \\z \\c "Figure"'  # "Table" of list of table and "Figure" for list of figure
    fldChar2 = OE('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    fldChar3 = OE('w:t')
    fldChar3.text = "Right-click to update field."
    fldChar2.append(fldChar3)

    fldChar4 = OE('w:fldChar')
    fldChar4.set(qn('w:fldCharType'), 'end')

    run._r.append(fldChar)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run._r.append(fldChar4)
def heading_to_index():
    paragraph = document.add_paragraph()
    run = paragraph.add_run()
    fldChar = OxmlElement('w:fldChar')  # creates a new element
    fldChar.set(qn('w:fldCharType'), 'begin')  # sets attribute on element
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')  # sets attribute on element
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u'  # change 1-3 depending on heading levels you need

    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    fldChar3 = OxmlElement('w:updateFields')
    fldChar3.set(qn('w:val'), 'true')
    # fldChar3.text = "Right-click to update field."
    fldChar2.append(fldChar3)

    fldChar4 = OxmlElement('w:fldChar')
    fldChar4.set(qn('w:fldCharType'), 'end')

    r_element = run._r
    r_element.append(fldChar)
    r_element.append(instrText)
    r_element.append(fldChar2)
    r_element.append(fldChar4)
    p_element = paragraph._p
def write_paragraph_new_style(content, font_name="Times New Roman", font_size=12, font_bold=False, font_italic=False,
                              font_underline=False, color=RGBColor(0, 0, 0), align="left", bef_spacing=5, aft_spacing=5,
                              line_spacing=1.5, set_style_name="my_style_namee",
                              keep_together=True, keep_with_next=False, page_break_before=False, widow_control=False):
    # paragraph = document.add_paragraph(content, style=use_style)

    # Create a style for my paragraph
    paragraph = document.add_paragraph(content)
    paragraph.style = document.styles.add_style(set_style_name, WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = font_name
    font.size = Pt(font_size)
    font.bold = font_bold
    font.italic = font_italic
    font.underline = font_underline
    font.color.rgb = color
    # Line spacing.
    paragraph_format = paragraph.paragraph_format
    paragraph_format.space_before = Pt(bef_spacing)
    paragraph_format.space_after = Pt(aft_spacing)
    paragraph_format.line_spacing = line_spacing
    paragraph_format.keep_together = keep_together
    paragraph_format.keep_with_next = keep_with_next
    paragraph_format.page_break_before = page_break_before
    paragraph_format.widow_control = widow_control
    if align.lower() == "left":
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    elif align.lower() == "right":
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    elif align.lower() == "center":
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    elif align.lower() == "justify":
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    else:
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    return paragraph
def write_paragraph_existing_style(content, use_style=None,font_name="", font_size=0, font_bold="", font_italic="",
                              font_underline="", color="", align=""):

    paragraph = document.add_paragraph(content, style=use_style)
    font = paragraph.style.font
    if font_name == "":
        pass
    else:
        font.name = font_name
    if font_size == 0:
        pass
    else:
        font.size = Pt(font_size)
    if font_bold == "":
        pass
    else:
        font.bold = font_bold
    if font_italic == "":
        pass
    else:
        font.italic = font_italic

    if font_underline == "":
        pass
    else:
        font.underline = font_underline

    if font_underline =="":
        pass
    else:
        font.color.rgb = color

    if align.lower() == "left":
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    elif align.lower() == "right":
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    elif align.lower() == "center":
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    elif align.lower() == "justify":
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    else:
        pass

    return paragraph

def create_element(name):
    return OxmlElement(name)
def create_attribute(element, name, value):
    element.set(qn(name), value)
def add_page_number(run):
    fldChar1 = create_element('w:fldChar')
    create_attribute(fldChar1, 'w:fldCharType', 'begin')

    instrText = create_element('w:instrText')
    create_attribute(instrText, 'xml:space', 'preserve')
    instrText.text = "PAGE"

    fldChar2 = create_element('w:fldChar')
    create_attribute(fldChar2, 'w:fldCharType', 'end')

    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
def MarkIndexEntry(entry, paragraph):
    run = paragraph.add_run()
    r = run._r
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'begin')
    r.append(fldChar)

    run = paragraph.add_run()
    r = run._r
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = ' XE "%s" ' % (entry)
    r.append(instrText)

    run = paragraph.add_run()
    r = run._r
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'end')
    r.append(fldChar)
def Figure(paragraph):
    run = run = paragraph.add_run()
    r = run._r
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'begin')
    r.append(fldChar)
    instrText = OxmlElement('w:instrText')
    instrText.text = ' SEQ Figure \* ARABIC'
    r.append(instrText)
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'end')
    r.append(fldChar)
def Table(paragraph):
    run = run = paragraph.add_run()
    r = run._r
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'begin')
    r.append(fldChar)
    instrText = OxmlElement('w:instrText')
    instrText.text = ' SEQ Table \* ARABIC'
    r.append(instrText)
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'end')
    r.append(fldChar)


# Adding Page Numbers in Footer (añadiendo numeros de pagina)
document.sections[0].footer.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
add_page_number(document.sections[0].footer.paragraphs[0].add_run())

# Controlling Margins (Control de Margenes de documento)
sections = document.sections
for section in sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

# dividing word into columns (research paper like) (dividir hojas en columna de ser necesario)
section = document.sections[0]
sectPr = section._sectPr
cols = sectPr.xpath("./w:cols")[0]
cols.set(qn("w:num"), "1")  # Number of columns, 1 is default.

# Main document. (Comienzo de principal documento)

#Portada
document.add_paragraph("")
pic = document.add_picture("img.png", width=Inches(5), height=Inches(2))
last_paragraph = document.paragraphs[-1]
last_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # Align Image to Center.
document.add_paragraph("")
document.add_paragraph("")
document.add_paragraph("")
p0 = write_paragraph_new_style(f"{titulo_port}", font_size=20, font_bold=True, set_style_name="bullshi",
                               color=RGBColor(0, 0, 0),
                               align="center", )
document.add_paragraph("")
document.add_paragraph("")
document.add_paragraph("")
p1 = write_paragraph_existing_style(f"Fecha: {fecha_hoy}", use_style="bullshi", align="center")

document.add_page_break()

#Indice de Contenidos
heading = document.add_heading("Índice de Contenidos", level=0)
heading_to_index()
document.add_page_break()

#Indice de Figuras
document.add_heading("Índice de Figuras", level=0)
add_list_of_figure()
document.add_page_break()

#Indice de Tablas
document.add_heading("Índice de Tablas", level=0)
add_list_of_table()
document.add_page_break()

# Giving headings that need to be included in Table of contents
document.add_heading("Trabajadores del día", level=1)
write_paragraph_existing_style(f"{trabajador_texto}")

document.add_page_break()
document.add_heading("Cabecera de Prueba", level=1)

document.add_paragraph('')
document.add_picture("img.png", width=Inches(5), height=Inches(2))
last_paragraph = document.paragraphs[-1]
last_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # Align Image to Center.
paragraph = write_paragraph_existing_style('Figura ', use_style='Caption',align="center", font_size=10, color=RGBColor(0,0,0))
Figure(paragraph)
paragraph.add_run(": Logo OGM")

document.add_paragraph('')


paragraph = write_paragraph_existing_style('Tabla ', use_style='Caption',align="center", font_size=12,color=RGBColor(0,0,0))
Table(paragraph)
paragraph.add_run(f": {nombre_tabla_1}")
df = pd.read_excel('ejemplo.xlsx',sheet_name='Sheet1')
crear_tabla(df)





#last_paragraph = document.paragraphs[-1]
#last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

# p0 = write_paragraph("Matias", font_size=20, font_bold=True, set_style_name="bullshit", color=RGBColor(255, 0, 0))

# p1 = write_paragraph("I am the king of the world", use_style="List Number")
# = write_paragraph("I am the king of the world", use_style="List Number")

# p3 = write_paragraph("God of the air", set_style_name="caca", color=RGBColor(0, 0, 255), font_size=13)
# p3.style.font.size = Pt(13) # this modifies all paragraphs with the same style, not only p3

# pic = document.add_picture("demo_pic.jpg", width=Inches(2.5), height=Inches(2.5))
# last_paragraph = document.paragraphs[-1]
# last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Align Image to Center.

document.save("word.docx")

word = win32com.client.DispatchEx("Word.Application")
doc = word.Documents.Open("C:/Users/56976/PycharmProjects/Automate_word/word.docx")
doc.TablesOfContents(1).Update()
doc.TablesOfFigures(1).Update()
doc.TablesOfFigures(2).Update()
doc.Close(SaveChanges=True)
word.Quit()
